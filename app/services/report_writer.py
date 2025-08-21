import io
from datetime import date
from typing import Dict, Any

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from app.schemas.report_auto import AutoReportData


class ReportWriter:
    """レポート生成サービス"""
    
    @staticmethod
    def generate_report_text(data: AutoReportData) -> str:
        """CSR/IR向け日本語文章を生成"""
        
        # 前年同期比の表現
        def get_change_text(change_percent: float, metric_name: str) -> str:
            if change_percent > 0:
                return f"{metric_name}は前年同期比{change_percent:.1f}%増加"
            elif change_percent < 0:
                return f"{metric_name}は前年同期比{abs(change_percent):.1f}%削減"
            else:
                return f"{metric_name}は前年同期比横ばい"
        
        text = f"""
環境への取り組み報告書
対象期間：{data.period_start.strftime('%Y年%m月%d日')} ～ {data.period_end.strftime('%Y年%m月%d日')}

## 1. 取り組みの概要

当社では、持続可能な経営の実現に向けて、全社一丸となった省エネルギー活動を推進しております。
従業員参加型のエネルギー管理システムを導入し、日常業務における環境負荷削減に取り組んでまいりました。

## 2. 実績・成果

### エネルギー使用実績
- 電力使用量：{data.energy_metrics.electricity_kwh:,.0f} kWh
- ガス使用量：{data.energy_metrics.gas_m3:,.0f} m³
- CO₂排出削減量：{data.energy_metrics.co2_reduction_kg:,.0f} kg

### 前年同期比較
- {get_change_text(data.comparison.electricity_change_percent, "電力使用量")}
- {get_change_text(data.comparison.gas_change_percent, "ガス使用量")}  
- {get_change_text(data.comparison.co2_change_percent, "CO₂排出量")}

### 従業員参加状況
全従業員{data.participation.total_employees:,}名のうち{data.participation.participating_employees:,}名が活動に参加し、
参加率{data.participation.participation_rate:.1f}%を達成いたしました。

## 3. 今後の展望

引き続き、従業員一人ひとりの環境意識向上を図り、より効果的な省エネルギー活動を推進してまいります。
デジタル技術を活用したエネルギー管理の高度化により、さらなる環境負荷削減を目指します。

これらの取り組みを通じて、持続可能な社会の実現に貢献し、企業価値の向上を図ってまいります。
"""
        return text.strip()
    
    @staticmethod
    def generate_pdf(data: AutoReportData) -> bytes:
        """PDFレポートを生成"""
        buffer = io.BytesIO()
        
        # A4サイズでPDF作成
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # スタイル設定
        styles = getSampleStyleSheet()
        
        # 日本語対応のため標準フォントを使用
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12
        )
        
        # コンテンツ構築
        story = []
        
        # タイトル
        story.append(Paragraph("環境への取り組み報告書", title_style))
        story.append(Paragraph(f"対象期間：{data.period_start.strftime('%Y年%m月%d日')} ～ {data.period_end.strftime('%Y年%m月%d日')}", normal_style))
        story.append(Spacer(1, 20))
        
        # 実績データテーブル
        story.append(Paragraph("実績サマリー", heading_style))
        
        table_data = [
            ['項目', '実績値', '前年同期比'],
            ['電力使用量', f'{data.energy_metrics.electricity_kwh:,.0f} kWh', f'{data.comparison.electricity_change_percent:+.1f}%'],
            ['ガス使用量', f'{data.energy_metrics.gas_m3:,.0f} m³', f'{data.comparison.gas_change_percent:+.1f}%'],
            ['CO₂削減量', f'{data.energy_metrics.co2_reduction_kg:,.0f} kg', f'{data.comparison.co2_change_percent:+.1f}%'],
            ['従業員参加率', f'{data.participation.participation_rate:.1f}%', '-']
        ]
        
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
        
        # 生成された文章を追加
        report_text = ReportWriter.generate_report_text(data)
        paragraphs = report_text.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                if para.startswith('##'):
                    # 見出し
                    heading_text = para.replace('##', '').strip()
                    story.append(Paragraph(heading_text, heading_style))
                elif para.startswith('#'):
                    # タイトル（既に追加済みなのでスキップ）
                    continue
                else:
                    # 通常の段落
                    story.append(Paragraph(para.strip(), normal_style))
        
        # PDF生成
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def generate_docx(data: AutoReportData) -> bytes:
        """Wordドキュメントを生成"""
        doc = Document()
        
        # タイトル
        title = doc.add_heading('環境への取り組み報告書', 0)
        title.alignment = 1  # Center
        
        # 期間
        period_para = doc.add_paragraph()
        period_para.add_run(f"対象期間：{data.period_start.strftime('%Y年%m月%d日')} ～ {data.period_end.strftime('%Y年%m月%d日')}")
        period_para.alignment = 1  # Center
        
        # 実績テーブル
        doc.add_heading('実績サマリー', level=1)
        
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # ヘッダー
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '項目'
        hdr_cells[1].text = '実績値'
        hdr_cells[2].text = '前年同期比'
        
        # データ
        table.rows[1].cells[0].text = '電力使用量'
        table.rows[1].cells[1].text = f'{data.energy_metrics.electricity_kwh:,.0f} kWh'
        table.rows[1].cells[2].text = f'{data.comparison.electricity_change_percent:+.1f}%'
        
        table.rows[2].cells[0].text = 'ガス使用量'
        table.rows[2].cells[1].text = f'{data.energy_metrics.gas_m3:,.0f} m³'
        table.rows[2].cells[2].text = f'{data.comparison.gas_change_percent:+.1f}%'
        
        table.rows[3].cells[0].text = 'CO₂削減量'
        table.rows[3].cells[1].text = f'{data.energy_metrics.co2_reduction_kg:,.0f} kg'
        table.rows[3].cells[2].text = f'{data.comparison.co2_change_percent:+.1f}%'
        
        table.rows[4].cells[0].text = '従業員参加率'
        table.rows[4].cells[1].text = f'{data.participation.participation_rate:.1f}%'
        table.rows[4].cells[2].text = '-'
        
        # 生成された文章を追加
        report_text = ReportWriter.generate_report_text(data)
        paragraphs = report_text.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                if para.startswith('##'):
                    # 見出し
                    heading_text = para.replace('##', '').strip()
                    doc.add_heading(heading_text, level=1)
                elif para.startswith('#'):
                    # タイトル（既に追加済みなのでスキップ）
                    continue
                else:
                    # 通常の段落
                    doc.add_paragraph(para.strip())
        
        # バイト形式で返却
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()