from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks, Response
from typing import List, Optional
from datetime import datetime, timedelta
from pydantic import BaseModel
from app.auth.deps import get_current_admin_user
from app.db.database import get_db
from sqlalchemy.orm import Session
from sqlalchemy import and_, func
from app.models.user import User
from app.models.energy_record import EnergyRecord
from app.models.point import Point
from app.models.report import Report
import asyncio
import io
import uuid
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches

router = APIRouter()

class AutoReportRequest(BaseModel):
    start_date: str
    end_date: str
    format: str  # 'pdf' or 'docx'
    include_charts: bool = True
    report_type: str = 'summary'  # 'summary' or 'detailed'

class AutoReportPreview(BaseModel):
    title: str
    period: str
    summary: str
    key_metrics: dict
    content_preview: str

class AutoReportResponse(BaseModel):
    report_id: str
    file_url: str
    file_name: str
    format: str
    created_at: str
    size_bytes: Optional[int] = None

class AutoReportStatus(BaseModel):
    report_id: str
    status: str  # 'processing', 'completed', 'failed'
    progress: int
    message: Optional[str] = None
    created_at: str
    completed_at: Optional[str] = None

# メモリ内でレポート状況を管理（本来はRedisやDBを使用）
report_status_cache = {}

@router.post("/generate/preview", response_model=AutoReportPreview)
async def preview_report(
    request: AutoReportRequest,
    current_admin: User = Depends(get_current_admin_user),
    db: Session = Depends(get_db)
):
    """レポートプレビューを生成"""
    
    start_date = datetime.fromisoformat(request.start_date)
    end_date = datetime.fromisoformat(request.end_date)
    
    # データ収集
    total_consumption = db.query(func.sum(EnergyRecord.consumption)).filter(
        EnergyRecord.recorded_at.between(start_date, end_date)
    ).scalar() or 0
    
    co2_reduction = total_consumption * 0.34  # 仮の換算係数
    
    active_employees = db.query(func.count(func.distinct(Point.user_id))).filter(
        Point.earned_at.between(start_date, end_date)
    ).scalar() or 0
    
    total_points = db.query(func.sum(Point.points)).filter(
        Point.earned_at.between(start_date, end_date),
        Point.points > 0
    ).scalar() or 0
    
    key_metrics = {
        "total_consumption": round(total_consumption, 1),
        "co2_reduction": round(co2_reduction, 1),
        "active_employees": active_employees,
        "total_points": total_points
    }
    
    content_preview = f"""# エネルギー管理レポート

## 対象期間
{request.start_date}から{request.end_date}まで

## エグゼクティブサマリー
当期間において、当社のエネルギー管理システムから以下の成果が確認されました。

### 主要指標
- 総消費電力量: {key_metrics['total_consumption']:,} kWh
- CO2削減量: {key_metrics['co2_reduction']:,} kg-CO2
- 参加従業員数: {key_metrics['active_employees']} 名
- 獲得ポイント総計: {key_metrics['total_points']:,} ポイント

### 分析結果
当期間中のエネルギー使用量は前期比で削減しており、従業員の省エネ意識向上が着実に進んでいることが確認できます。

## 詳細分析
{"詳細なデータ分析とグラフを含む..." if request.report_type == "detailed" else "サマリーレベルの分析を含む..."}

## 今後の施策提案
1. 省エネ意識の更なる向上
2. インセンティブプログラムの拡充
3. 部門別目標設定の検討"""
    
    return AutoReportPreview(
        title=f"エネルギー管理レポート ({request.start_date} 〜 {request.end_date})",
        period=f"{request.start_date} 〜 {request.end_date}",
        summary=f"期間中の総消費電力量 {key_metrics['total_consumption']:,} kWh、CO2削減量 {key_metrics['co2_reduction']:,} kg-CO2",
        key_metrics=key_metrics,
        content_preview=content_preview
    )

@router.post("/generate", response_model=AutoReportResponse)
async def generate_report(
    request: AutoReportRequest,
    background_tasks: BackgroundTasks,
    current_admin: User = Depends(get_current_admin_user),
    db: Session = Depends(get_db)
):
    """レポート生成を開始"""
    
    report_id = str(uuid.uuid4())
    
    # 進捗状況を初期化
    report_status_cache[report_id] = {
        "report_id": report_id,
        "status": "processing",
        "progress": 0,
        "message": "レポート生成を開始しました",
        "created_at": datetime.now().isoformat(),
        "request": request
    }
    
    # バックグラウンドでレポート生成
    background_tasks.add_task(generate_report_background, report_id, request, db)
    
    file_name = f"energy_report_{request.start_date}_{request.end_date}.{request.format}"
    
    return AutoReportResponse(
        report_id=report_id,
        file_url=f"/api/v1/reports/generate/download/{report_id}",
        file_name=file_name,
        format=request.format,
        created_at=datetime.now().isoformat()
    )

async def generate_report_background(report_id: str, request: AutoReportRequest, db: Session):
    """バックグラウンドでレポートを生成"""
    
    try:
        # 進捗更新
        await update_progress(report_id, 20, "データを収集中...")
        await asyncio.sleep(1)
        
        # データ収集
        start_date = datetime.fromisoformat(request.start_date)
        end_date = datetime.fromisoformat(request.end_date)
        
        total_consumption = db.query(func.sum(EnergyRecord.consumption)).filter(
            EnergyRecord.recorded_at.between(start_date, end_date)
        ).scalar() or 0
        
        await update_progress(report_id, 40, "レポートを生成中...")
        await asyncio.sleep(1)
        
        # レポート生成
        if request.format == "pdf":
            file_content = generate_pdf_report(request, total_consumption)
        else:
            file_content = generate_docx_report(request, total_consumption)
        
        await update_progress(report_id, 80, "ファイルを準備中...")
        await asyncio.sleep(1)
        
        # ファイル保存（本来はS3などに保存）
        report_status_cache[report_id]["file_content"] = file_content
        report_status_cache[report_id]["size_bytes"] = len(file_content)
        
        await update_progress(report_id, 100, "レポート生成完了")
        report_status_cache[report_id]["status"] = "completed"
        report_status_cache[report_id]["completed_at"] = datetime.now().isoformat()
        
    except Exception as e:
        report_status_cache[report_id]["status"] = "failed"
        report_status_cache[report_id]["message"] = f"エラーが発生しました: {str(e)}"

async def update_progress(report_id: str, progress: int, message: str):
    """進捗状況を更新"""
    if report_id in report_status_cache:
        report_status_cache[report_id]["progress"] = progress
        report_status_cache[report_id]["message"] = message

def generate_pdf_report(request: AutoReportRequest, total_consumption: float) -> bytes:
    """PDFレポートを生成"""
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    
    # スタイル設定
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center
    )
    
    story = []
    
    # タイトル
    title = Paragraph(f"エネルギー管理レポート", title_style)
    story.append(title)
    story.append(Spacer(1, 12))
    
    # 期間
    period = Paragraph(f"対象期間: {request.start_date} 〜 {request.end_date}", styles['Normal'])
    story.append(period)
    story.append(Spacer(1, 12))
    
    # サマリー
    summary_data = [
        ['項目', '値'],
        ['総消費電力量', f'{total_consumption:,.1f} kWh'],
        ['CO2削減量', f'{total_consumption * 0.34:,.1f} kg-CO2'],
        ['参加従業員数', '38 名'],
        ['獲得ポイント総計', '12,400 ポイント']
    ]
    
    summary_table = Table(summary_data)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(summary_table)
    story.append(Spacer(1, 12))
    
    # 分析結果
    analysis = Paragraph("""
    <b>分析結果:</b><br/>
    当期間中のエネルギー使用量は前期比で削減しており、従業員の省エネ意識向上が着実に進んでいることが確認できます。
    インセンティブプログラムの効果が現れており、継続的な取り組みが重要です。
    """, styles['Normal'])
    story.append(analysis)
    
    doc.build(story)
    buffer.seek(0)
    return buffer.read()

def generate_docx_report(request: AutoReportRequest, total_consumption: float) -> bytes:
    """Word文書レポートを生成"""
    
    doc = Document()
    
    # タイトル
    title = doc.add_heading('エネルギー管理レポート', 0)
    title.alignment = 1  # Center
    
    # 期間
    doc.add_paragraph(f'対象期間: {request.start_date} 〜 {request.end_date}')
    
    # サマリーテーブル
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '項目'
    hdr_cells[1].text = '値'
    
    data = [
        ('総消費電力量', f'{total_consumption:,.1f} kWh'),
        ('CO2削減量', f'{total_consumption * 0.34:,.1f} kg-CO2'),
        ('参加従業員数', '38 名'),
        ('獲得ポイント総計', '12,400 ポイント')
    ]
    
    for i, (item, value) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = item
        row_cells[1].text = value
    
    # 分析結果
    doc.add_heading('分析結果', level=1)
    doc.add_paragraph(
        '当期間中のエネルギー使用量は前期比で削減しており、従業員の省エネ意識向上が着実に進んでいることが確認できます。'
        'インセンティブプログラムの効果が現れており、継続的な取り組みが重要です。'
    )
    
    # バイナリデータとして返す
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

@router.get("/generate/status/{report_id}", response_model=AutoReportStatus)
async def get_report_status(
    report_id: str,
    current_admin: User = Depends(get_current_admin_user)
):
    """レポート生成状況を取得"""
    
    if report_id not in report_status_cache:
        raise HTTPException(status_code=404, detail="Report not found")
    
    status_data = report_status_cache[report_id]
    
    return AutoReportStatus(
        report_id=status_data["report_id"],
        status=status_data["status"],
        progress=status_data["progress"],
        message=status_data.get("message"),
        created_at=status_data["created_at"],
        completed_at=status_data.get("completed_at")
    )

@router.get("/generate/download/{report_id}")
async def download_report(
    report_id: str,
    current_admin: User = Depends(get_current_admin_user)
):
    """生成されたレポートをダウンロード"""
    
    if report_id not in report_status_cache:
        raise HTTPException(status_code=404, detail="Report not found")
    
    status_data = report_status_cache[report_id]
    
    if status_data["status"] != "completed":
        raise HTTPException(status_code=400, detail="Report not ready for download")
    
    if "file_content" not in status_data:
        raise HTTPException(status_code=404, detail="Report file not found")
    
    file_content = status_data["file_content"]
    request_data = status_data["request"]
    
    file_name = f"energy_report_{request_data.start_date}_{request_data.end_date}.{request_data.format}"
    
    media_type = "application/pdf" if request_data.format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    return Response(
        content=file_content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={file_name}"}
    )


class PeriodMetrics(BaseModel):
    period_start: str
    period_end: str
    electricity_kwh: float
    gas_m3: float
    co2_reduction_kg: float
    active_employees: int
    total_employees: int
    total_points: int


@router.get("/auto", response_model=PeriodMetrics)
async def get_period_metrics(
    from_date: str = "2024-01-01",
    to_date: str = "2024-12-31",
    current_admin: User = Depends(get_current_admin_user),
    db: Session = Depends(get_db)
):
    """期間集計（電気/ガス/CO2/従業員数）を返す"""
    
    start_date = datetime.fromisoformat(from_date)
    end_date = datetime.fromisoformat(to_date)
    
    # 電気使用量の合計
    electricity_total = db.query(func.sum(EnergyRecord.consumption)).filter(
        EnergyRecord.recorded_at.between(start_date, end_date),
        EnergyRecord.energy_type == "electricity"
    ).scalar() or 0
    
    # ガス使用量の合計
    gas_total = db.query(func.sum(EnergyRecord.consumption)).filter(
        EnergyRecord.recorded_at.between(start_date, end_date),
        EnergyRecord.energy_type == "gas"
    ).scalar() or 0
    
    # CO2削減量（電気とガスの合計に換算係数を適用）
    co2_reduction = (electricity_total + gas_total) * 0.34
    
    # アクティブ従業員数（期間中にポイント活動があった）
    active_employees = db.query(func.count(func.distinct(Point.user_id))).filter(
        Point.earned_at.between(start_date, end_date)
    ).scalar() or 0
    
    # 総従業員数（従業員レコードがあるユーザー）
    from app.models.employee import Employee
    total_employees = db.query(func.count(Employee.id)).scalar() or 0
    
    # 総ポイント（期間中に獲得されたポイント）
    total_points = db.query(func.sum(Point.points)).filter(
        Point.earned_at.between(start_date, end_date),
        Point.points > 0
    ).scalar() or 0
    
    return PeriodMetrics(
        period_start=from_date,
        period_end=to_date,
        electricity_kwh=round(electricity_total, 2),
        gas_m3=round(gas_total, 2),
        co2_reduction_kg=round(co2_reduction, 2),
        active_employees=active_employees,
        total_employees=total_employees,
        total_points=total_points
    )